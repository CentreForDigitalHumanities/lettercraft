from io import TextIOWrapper
from typing import Dict

from django.utils.html import strip_tags

from docx import Document
from docx.document import Document as DocumentObject


def save_docx(data: Dict, out: TextIOWrapper) -> None:
    document = Document()
    _add_preamble(data["metadata"], document)
    for source in data['sources']:
        _add_source(source, document)

    document.save(out)


def _add_preamble(data: Dict, document: DocumentObject) -> None:
    document.add_heading(f"Lettercraft data")
    document.add_paragraph(
        f"Exported from {data["url"]} on {data["date"]}"
    )


def _add_source(data: Dict, document: DocumentObject) -> None:
    document.add_page_break()
    document.add_heading(data["name"])

    if reference := data['reference']:
        document.add_paragraph(strip_tags(reference))

    if contributors := data["contributors"]:
        _add_key_value_paragraph("contributors", ", ".join(contributors), document)

    if description := data["description"]:
        document.add_paragraph(strip_tags(description))

    document.add_heading("Episodes", level=2)
    for episode in data["episodes"]:
        _add_episode(episode, data, document)

    document.add_heading("Agents", level=2)
    for agent in data["agents"]:
        _add_agent(agent, document)

    for key in ["letters", "gifts", "locations"]:
        document.add_heading(key.capitalize(), level=2)
        for entity in data[key]:
            _add_entity(entity, document)


def _add_episode(data: Dict, source_data: Dict, document: DocumentObject) -> None:
    document.add_heading(data["name"], level=3)

    if source_ref := _episode_source_ref(data["source_reference"]):
        p = document.add_paragraph()
        run = p.add_run(source_ref)
        run.italic = True

    if summary := data["summary"]:
        document.add_paragraph(summary)

    if labels := data["labels"]:
        _add_key_value_paragraph("labels", ", ".join(labels), document)

    if designators := data["designators"]:
        _add_key_value_paragraph("designators", ", ".join(designators), document)

    for key in ["agents", "letters", "gifts", "locations"]:
        if entities := _episode_entity_names(data, source_data, key):
            _add_key_value_paragraph(key, entities, document)


def _add_key_value_paragraph(key: str, value: str, document: DocumentObject):
    p = document.add_paragraph()
    head = p.add_run(key.capitalize() + ":")
    head.bold = True
    p.add_run(" " + value)


def _episode_source_ref(data: Dict) -> str:
    keys = ["book", "chapter", "page"]
    parts = [
        f"{key} {data[key]}"
        for key in keys
        if data[key] is not None
    ]
    return ", ".join(parts).capitalize()


def _episode_entity_names(episode_data: Dict, source_data: Dict, key: str):
    ids = episode_data[key]
    all_entities = source_data[key]
    entities = [
        next(entity for entity in all_entities if entity["id"] == entity_id)
        for entity_id in ids
    ]
    names = [entity["name"] for entity in entities]
    return ", ".join(names)


def _add_agent(data: Dict, document: DocumentObject):
    p = document.add_paragraph(style="List Bullet")
    name = p.add_run(data["name"])
    if data["is_group"]:
        name.add_text(" [group]")
    name.add_text(". ")

    description = p.add_run(data["description"])
    description.italic = True


def _add_entity(data: Dict, document: DocumentObject):
    p = document.add_paragraph(style="List Bullet")
    p.add_run(data["name"] + ". ")

    description = p.add_run(data["description"])
    description.italic = True
